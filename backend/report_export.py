"""
Generare rapoarte (PDF / XLSX / DOCX) cu antet personalizabil.
- PDF: imaginea de antet este desenata ca FUNDAL pe toata pagina A4, in spatele textului.
- DOCX: imaginea de antet este plasata ca watermark (behindDoc) pe toata pagina A4 in fiecare sectiune.
- XLSX: fara antet (cerinta utilizator).
Antet implicit: backend/assets/header_default.jpg
Antet custom (upload): backend/assets/header_custom.* (jpg/png)
"""
import io
import os
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional

ASSETS_DIR = Path(__file__).parent / "assets"
ASSETS_DIR.mkdir(exist_ok=True)
DEFAULT_HEADER = ASSETS_DIR / "header_default.jpg"

COLUMN_LABELS = {
    "company_name": "Companie",
    "platform": "Platforma",
    "vm_name": "VM / Masina",
    "status": "Stare",
    "backup_date": "Data",
    "size": "Dimensiune",
    "transferred": "Transferat",
    "duration": "Durata",
    "details": "Detalii",
    "source": "Sursa",
    "from_address": "Expeditor",
    "email_subject": "Subiect Email",
    "email_date": "Data Email",
    "is_unknown": "Neidentificat",
}

STATUS_LABELS = {"success": "Succes", "failed": "Esuat", "warning": "Avertisment"}


def get_header_path() -> Optional[str]:
    for ext in ("jpg", "jpeg", "png"):
        custom = ASSETS_DIR / f"header_custom.{ext}"
        if custom.exists():
            return str(custom)
    if DEFAULT_HEADER.exists():
        return str(DEFAULT_HEADER)
    return None


def save_custom_header(content: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext not in ("jpg", "jpeg", "png"):
        raise ValueError("Format antet invalid (doar jpg/png)")
    for e in ("jpg", "jpeg", "png"):
        p = ASSETS_DIR / f"header_custom.{e}"
        if p.exists():
            p.unlink()
    path = ASSETS_DIR / f"header_custom.{ext}"
    path.write_bytes(content)
    return str(path)


def reset_header():
    for e in ("jpg", "jpeg", "png"):
        p = ASSETS_DIR / f"header_custom.{e}"
        if p.exists():
            p.unlink()


def _format_value(backup: Dict[str, Any], col: str) -> str:
    v = backup.get(col, "")
    if col == "status":
        return STATUS_LABELS.get(v, v or "")
    if col == "is_unknown":
        return "Da" if v else ""
    if v is None:
        return ""
    return str(v)


def _build_title(scope: str, company_name: str, date_from: str, date_to: str) -> str:
    if scope == "company" and company_name:
        base = f"Raport Backup – {company_name}"
    else:
        base = "Raport Backup – Toate Job-urile"
    if date_from and date_to:
        if date_from == date_to:
            base += f" ({date_from})"
        else:
            base += f" ({date_from} – {date_to})"
    elif date_from:
        base += f" (de la {date_from})"
    elif date_to:
        base += f" (pana la {date_to})"
    return base


# ─── PDF ───
def generate_pdf(backups: List[Dict], columns: List[str], scope: str, company_name: str,
                 date_from: str, date_to: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=3.5 * cm, bottomMargin=1.5 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Title"], fontSize=16, spaceAfter=12, alignment=1)
    sub_style = ParagraphStyle("sub", parent=styles["Normal"], fontSize=9, textColor=colors.grey, alignment=1, spaceAfter=10)
    cell_style = ParagraphStyle("cell", parent=styles["Normal"], fontSize=8, leading=10)
    header_cell_style = ParagraphStyle("hcell", parent=styles["Normal"], fontSize=8, leading=10, textColor=colors.white, alignment=1)

    flow = []

    flow.append(Paragraph(_build_title(scope, company_name, date_from, date_to), title_style))
    flow.append(Paragraph(f"Generat: {datetime.now().strftime('%d.%m.%Y %H:%M')} &middot; Total job-uri: {len(backups)}", sub_style))

    total = len(backups)
    n_ok = sum(1 for b in backups if b.get("status") == "success")
    n_fail = sum(1 for b in backups if b.get("status") == "failed")
    n_warn = sum(1 for b in backups if b.get("status") == "warning")
    summary = [["Total", "Succes", "Esuat", "Avertisment"], [str(total), str(n_ok), str(n_fail), str(n_warn)]]
    s_table = Table(summary, colWidths=[3.5 * cm] * 4)
    s_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("TEXTCOLOR", (1, 1), (1, 1), colors.HexColor("#059669")),
        ("TEXTCOLOR", (2, 1), (2, 1), colors.HexColor("#dc2626")),
        ("TEXTCOLOR", (3, 1), (3, 1), colors.HexColor("#d97706")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    flow.append(s_table)
    flow.append(Spacer(1, 0.5 * cm))

    if not columns:
        columns = ["company_name", "platform", "vm_name", "status", "backup_date", "size", "duration"]

    head_row = [Paragraph(COLUMN_LABELS.get(c, c), header_cell_style) for c in columns]
    data = [head_row]
    for b in backups:
        row = [Paragraph(_format_value(b, c)[:200], cell_style) for c in columns]
        data.append(row)

    page_w = A4[0] - 3 * cm
    col_w = page_w / len(columns)
    table = Table(data, colWidths=[col_w] * len(columns), repeatRows=1)
    style_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#374151")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#cbd5e1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]
    if "status" in columns:
        s_idx = columns.index("status")
        for i, b in enumerate(backups, start=1):
            status = b.get("status")
            if status == "success":
                style_cmds.append(("TEXTCOLOR", (s_idx, i), (s_idx, i), colors.HexColor("#059669")))
            elif status == "failed":
                style_cmds.append(("TEXTCOLOR", (s_idx, i), (s_idx, i), colors.HexColor("#dc2626")))
            elif status == "warning":
                style_cmds.append(("TEXTCOLOR", (s_idx, i), (s_idx, i), colors.HexColor("#d97706")))
    table.setStyle(TableStyle(style_cmds))
    flow.append(table)

    # Background image on every page (full A4, behind content)
    header_path = get_header_path()

    def _draw_background(canvas, _doc):
        if not header_path:
            return
        try:
            canvas.saveState()
            page_w_pt, page_h_pt = A4
            canvas.drawImage(
                header_path, 0, 0,
                width=page_w_pt, height=page_h_pt,
                preserveAspectRatio=False, mask='auto'
            )
            canvas.restoreState()
        except Exception:
            pass

    doc.build(flow, onFirstPage=_draw_background, onLaterPages=_draw_background)
    return buf.getvalue()


# ─── XLSX (fara antet) ───
def generate_xlsx(backups: List[Dict], columns: List[str], scope: str, company_name: str,
                  date_from: str, date_to: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Raport Backup"

    next_row = 1
    ws.cell(row=next_row, column=1, value=_build_title(scope, company_name, date_from, date_to))
    ws.cell(row=next_row, column=1).font = Font(name="Calibri", size=14, bold=True)
    ws.merge_cells(start_row=next_row, end_row=next_row, start_column=1, end_column=max(len(columns) if columns else 4, 4))
    next_row += 1
    ws.cell(row=next_row, column=1, value=f"Generat: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  Total: {len(backups)} job-uri")
    ws.cell(row=next_row, column=1).font = Font(name="Calibri", size=10, italic=True, color="666666")
    ws.merge_cells(start_row=next_row, end_row=next_row, start_column=1, end_column=max(len(columns) if columns else 4, 4))
    next_row += 2

    if not columns:
        columns = ["company_name", "platform", "vm_name", "status", "backup_date", "size", "duration"]

    header_fill = PatternFill("solid", fgColor="374151")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    border = Border(
        left=Side(style="thin", color="CBD5E1"), right=Side(style="thin", color="CBD5E1"),
        top=Side(style="thin", color="CBD5E1"), bottom=Side(style="thin", color="CBD5E1"),
    )
    for c_idx, col in enumerate(columns, start=1):
        cell = ws.cell(row=next_row, column=c_idx, value=COLUMN_LABELS.get(col, col))
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
    header_row = next_row
    next_row += 1

    for b in backups:
        for c_idx, col in enumerate(columns, start=1):
            cell = ws.cell(row=next_row, column=c_idx, value=_format_value(b, col))
            cell.font = Font(name="Calibri", size=10)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = border
            if col == "status":
                if b.get("status") == "success":
                    cell.font = Font(name="Calibri", size=10, color="059669", bold=True)
                elif b.get("status") == "failed":
                    cell.font = Font(name="Calibri", size=10, color="DC2626", bold=True)
                elif b.get("status") == "warning":
                    cell.font = Font(name="Calibri", size=10, color="D97706", bold=True)
        next_row += 1

    widths = {"company_name": 24, "platform": 14, "vm_name": 22, "status": 14, "backup_date": 14,
              "size": 14, "transferred": 14, "duration": 12, "details": 50, "source": 12,
              "from_address": 28, "email_subject": 36, "email_date": 18, "is_unknown": 14}
    for c_idx, col in enumerate(columns, start=1):
        ws.column_dimensions[get_column_letter(c_idx)].width = widths.get(col, 18)

    ws.freeze_panes = ws.cell(row=header_row + 1, column=1)

    ws2 = wb.create_sheet("Sumar")
    n_ok = sum(1 for b in backups if b.get("status") == "success")
    n_fail = sum(1 for b in backups if b.get("status") == "failed")
    n_warn = sum(1 for b in backups if b.get("status") == "warning")
    summary_data = [
        ["Indicator", "Valoare"],
        ["Total job-uri", len(backups)],
        ["Succes", n_ok],
        ["Esuat", n_fail],
        ["Avertisment", n_warn],
        ["Perioada", f"{date_from or '-'} → {date_to or '-'}"],
        ["Scop", company_name if scope == "company" else "Toate companiile"],
    ]
    for r, row in enumerate(summary_data, start=1):
        for c, val in enumerate(row, start=1):
            cell = ws2.cell(row=r, column=c, value=val)
            if r == 1:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = header_fill
    ws2.column_dimensions["A"].width = 24
    ws2.column_dimensions["B"].width = 30

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ─── DOCX ───
def _add_docx_background(section, image_path: str):
    """
    Adauga imaginea ca watermark (behindDoc=1) pe toata pagina A4,
    plasata in header-ul sectiunii ca sa apara pe toate paginile.
    """
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    from docx.shared import Emu
    import hashlib

    header = section.header
    # asigura macar un paragraf in header
    if not header.paragraphs:
        header.add_paragraph()
    para = header.paragraphs[0]
    run = para.add_run()

    # adauga imaginea ca relatie in header part
    header_part = header.part
    with open(image_path, "rb") as f:
        img_bytes = f.read()
    ext = image_path.rsplit(".", 1)[-1].lower()
    if ext == "jpg":
        ext = "jpeg"
    from docx.image.image import Image as DocxImage
    image = DocxImage.from_blob(img_bytes)
    rel_id, image_part = header_part.get_or_add_image(image_path)

    # dimensiuni A4 in EMU (1 cm = 360000 EMU)
    page_w_emu = int(21.0 * 360000)   # 21 cm
    page_h_emu = int(29.7 * 360000)   # 29.7 cm

    pic_id = int(hashlib.md5(rel_id.encode()).hexdigest()[:6], 16) % 100000 + 1

    # XML pentru anchor behindDoc, pozitionat la coltul paginii, cu dimensiunea paginii
    drawing_xml = f'''
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
               xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="1"
                 behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
        <wp:extent cx="{page_w_emu}" cy="{page_h_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="{pic_id}" name="Background"/>
        <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="0"/></wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic>
              <pic:nvPicPr>
                <pic:cNvPr id="{pic_id}" name="Background"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip r:embed="{rel_id}"/>
                <a:stretch><a:fillRect/></a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{page_w_emu}" cy="{page_h_emu}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
    '''
    from docx.oxml import parse_xml
    drawing_el = parse_xml(drawing_xml)
    run._r.append(drawing_el)


def generate_docx(backups: List[Dict], columns: List[str], scope: str, company_name: str,
                  date_from: str, date_to: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.header_distance = Cm(0)
        section.footer_distance = Cm(0)

    # Antet ca background full-page A4 (in spatele textului), pe toate paginile
    header_path = get_header_path()
    if header_path:
        try:
            for section in document.sections:
                _add_docx_background(section, header_path)
        except Exception:
            pass

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(_build_title(scope, company_name, date_from, date_to))
    run.bold = True
    run.font.size = Pt(16)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"Generat: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  Total: {len(backups)} job-uri")
    sub_run.italic = True
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if not columns:
        columns = ["company_name", "platform", "vm_name", "status", "backup_date", "size", "duration"]

    n_ok = sum(1 for b in backups if b.get("status") == "success")
    n_fail = sum(1 for b in backups if b.get("status") == "failed")
    n_warn = sum(1 for b in backups if b.get("status") == "warning")
    sum_table = document.add_table(rows=2, cols=4)
    sum_table.style = "Light Grid Accent 1"
    sum_hdr = sum_table.rows[0].cells
    sum_val = sum_table.rows[1].cells
    for i, label in enumerate(["Total", "Succes", "Esuat", "Avertisment"]):
        sum_hdr[i].text = label
    for i, val in enumerate([str(len(backups)), str(n_ok), str(n_fail), str(n_warn)]):
        sum_val[i].text = val

    document.add_paragraph()

    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr[i].text = COLUMN_LABELS.get(col, col)
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True

    for b in backups:
        row = table.add_row().cells
        for i, col in enumerate(columns):
            row[i].text = _format_value(b, col)[:300]

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


# ─── CSV ───
def generate_csv(backups: List[Dict], columns: List[str]) -> bytes:
    import csv
    if not columns:
        columns = ["company_name", "platform", "vm_name", "status", "backup_date", "size", "duration"]
    buf = io.StringIO()
    writer = csv.writer(buf, delimiter=",", quoting=csv.QUOTE_MINIMAL)
    writer.writerow([COLUMN_LABELS.get(c, c) for c in columns])
    for b in backups:
        writer.writerow([_format_value(b, c) for c in columns])
    # BOM pentru Excel sa recunoasca UTF-8
    return ("\ufeff" + buf.getvalue()).encode("utf-8")
